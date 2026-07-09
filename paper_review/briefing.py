"""주간 브리핑 문서 생성 (.md + .docx).

출력 형식 규칙 반영:
- 지정 템플릿 순서: 제목·저자·게재일 → 핵심 기여 → 방법론 → 실무 적용 포인트 → 한계
- 각 요약에 원본 arXiv 링크와 논문 ID를 반드시 기재
- '실무 적용 포인트'는 에이전트 의견임을 라벨로 명시
- 관심도 순위 표 포함, [추출 실패]/[초록 기반 요약] 목록 분리
- 파일명 규칙: paper_brief_YYYYMMDD.docx (+ 동일 이름 .md)
- 문서 하단에 "사람 검토·승인 후 공유" 고지 (자동 발송·공유 금지 규칙)

제목·저자·게재일·링크는 전부 arXiv API 메타데이터(Paper)에서 렌더링한다.
LLM 출력(Summary)은 요약 본문 섹션에만 쓰인다.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from .arxiv_client import Paper
from .summarizer import Summary

APPROVAL_NOTICE = (
    "본 브리핑은 자동 생성된 초안입니다. 팀 공유·발송 전 반드시 사람이 내용을 "
    "검토·승인해야 합니다. '실무 적용 포인트'는 요약 에이전트의 의견이며 논문 "
    "내용과 구분됩니다."
)


@dataclass
class FailedItem:
    paper: Paper
    reason: str
    fallback_used: bool = False  # True 면 초록 기반 요약으로 대체됨


@dataclass
class BriefingData:
    period_label: str
    keywords: list[str]
    categories: list[str]
    items: list[tuple[Paper, Summary]] = field(default_factory=list)
    failed: list[FailedItem] = field(default_factory=list)
    skipped_duplicates: list[Paper] = field(default_factory=list)

    def ranked(self) -> list[tuple[Paper, Summary]]:
        return sorted(self.items, key=lambda t: t[1].interest_score, reverse=True)


def _fmt_date(iso: str) -> str:
    return iso[:10] if iso else "?"


def _fmt_authors(authors: list[str], limit: int = 6) -> str:
    if len(authors) <= limit:
        return ", ".join(authors)
    return ", ".join(authors[:limit]) + f" 외 {len(authors) - limit}명"


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------
def render_markdown(data: BriefingData) -> str:
    L: list[str] = []
    L.append(f"# 🔒 신규 정보보호 논문 주간 브리핑 ({data.period_label})")
    L.append("")
    L.append(f"- 수집 카테고리: {', '.join(data.categories)}")
    L.append(f"- 관심 키워드: {', '.join(data.keywords)}")
    L.append(f"- 요약 논문 수: {len(data.items)}편"
             f" / 추출 실패 {len(data.failed)}건 / 중복 제외 {len(data.skipped_duplicates)}건")
    L.append("")

    ranked = data.ranked()

    L.append("## 관심도 순위")
    L.append("")
    if ranked:
        L.append("| 순위 | 논문 (arXiv ID) | 게재일 | 관심도 | 근거 | 요약 근거 |")
        L.append("|---|---|---|---|---|---|")
        for rank, (p, s) in enumerate(ranked, 1):
            L.append(
                f"| {rank} | [{p.title}]({p.abs_url}) ({p.arxiv_id}) "
                f"| {_fmt_date(p.published)} | {s.interest_score}/5 "
                f"| {s.interest_reason} | {s.basis} |"
            )
    else:
        L.append("이번 주기에는 요약된 논문이 없습니다.")
    L.append("")

    L.append("## 논문별 요약")
    L.append("")
    for rank, (p, s) in enumerate(ranked, 1):
        # 템플릿 순서: 제목·저자·게재일 → 핵심 기여 → 방법론 → 실무 적용 포인트 → 한계
        L.append(f"### {rank}. {p.title}")
        L.append("")
        L.append(f"- **arXiv ID**: {p.arxiv_id} · **원문**: {p.abs_url}")
        L.append(f"- **저자**: {_fmt_authors(p.authors)}")
        L.append(f"- **게재일**: {_fmt_date(p.published)} · **카테고리**: {', '.join(p.categories)}")
        L.append(f"- **요약 근거**: {s.basis}")
        if s.warnings:
            for w in s.warnings:
                L.append(f"- ⚠️ {w}")
        L.append("")
        L.append("**핵심 기여**")
        L.append("")
        for c in s.key_contributions or ["(내용 없음)"]:
            L.append(f"- {c}")
        L.append("")
        L.append("**방법론**")
        L.append("")
        L.append(s.methodology or "(내용 없음)")
        L.append("")
        L.append("**실무 적용 포인트** *(에이전트 의견 — 논문 내용과 구분됨)*")
        L.append("")
        for pt in s.practical_points or ["(내용 없음)"]:
            L.append(f"- {pt}")
        L.append("")
        L.append("**한계**")
        L.append("")
        for lim in s.limitations or ["- 논문에 명시된 한계가 확인되지 않음"]:
            L.append(f"- {lim}" if not lim.startswith("-") else lim)
        L.append("")

    L.append("## [추출 실패] 목록")
    L.append("")
    if data.failed:
        for item in data.failed:
            note = " → 초록 기반 요약으로 대체" if item.fallback_used else " → 요약 제외"
            L.append(
                f"- {item.paper.title} ({item.paper.arxiv_id}) — {item.reason}{note}"
                f" · {item.paper.abs_url}"
            )
    else:
        L.append("- 없음")
    L.append("")

    if data.skipped_duplicates:
        L.append("## 중복 제외 (아카이브에 기존 요약 존재)")
        L.append("")
        for p in data.skipped_duplicates:
            L.append(f"- {p.title} ({p.arxiv_id})")
        L.append("")

    L.append("---")
    L.append("")
    L.append(f"> {APPROVAL_NOTICE}")
    L.append("")
    return "\n".join(L)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
# 한글 표시용 폰트. python-docx 기본 템플릿은 라틴 전용 테마 폰트를 쓰기 때문에
# 명시적으로 지정하지 않으면 뷰어에 따라 한글이 깨져(□) 보인다.
KOREAN_FONT = "맑은 고딕"  # Malgun Gothic — 미설치 환경에서는 Word가 자동 대체

_STYLES_TO_FIX = ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]


def _apply_korean_fonts(doc) -> None:
    """문서 스타일에 한글 폰트를 명시한다.

    - w:ascii / w:hAnsi / w:eastAsia / w:cs 를 모두 지정하고,
    - 테마 폰트 속성(asciiTheme 등)을 제거한다. (테마 속성이 남아 있으면
      Word 가 명시 폰트보다 테마를 우선해 한글 깨짐이 재발할 수 있다.)
    """
    from docx.oxml.ns import qn

    for name in _STYLES_TO_FIX:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = KOREAN_FONT  # w:ascii, w:hAnsi
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rfonts.attrib.pop(qn(f"w:{theme_attr}"), None)
        rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
        rfonts.set(qn("w:cs"), KOREAN_FONT)


def render_docx(data: BriefingData, dest: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _apply_korean_fonts(doc)
    doc.add_heading(f"신규 정보보호 논문 주간 브리핑 ({data.period_label})", level=0)

    meta = doc.add_paragraph()
    meta.add_run(
        f"수집 카테고리: {', '.join(data.categories)}\n"
        f"관심 키워드: {', '.join(data.keywords)}\n"
        f"요약 {len(data.items)}편 · 추출 실패 {len(data.failed)}건 · "
        f"중복 제외 {len(data.skipped_duplicates)}건"
    )

    ranked = data.ranked()

    doc.add_heading("관심도 순위", level=1)
    if ranked:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["순위", "논문 (arXiv ID)", "게재일", "관심도", "근거"]):
            hdr[i].text = h
        for rank, (p, s) in enumerate(ranked, 1):
            row = table.add_row().cells
            row[0].text = str(rank)
            row[1].text = f"{p.title} ({p.arxiv_id})"
            row[2].text = _fmt_date(p.published)
            row[3].text = f"{s.interest_score}/5"
            row[4].text = s.interest_reason
    else:
        doc.add_paragraph("이번 주기에는 요약된 논문이 없습니다.")

    doc.add_heading("논문별 요약", level=1)
    for rank, (p, s) in enumerate(ranked, 1):
        doc.add_heading(f"{rank}. {p.title}", level=2)
        info = doc.add_paragraph()
        info.add_run(
            f"arXiv ID: {p.arxiv_id}\n"
            f"원문: {p.abs_url}\n"
            f"저자: {_fmt_authors(p.authors)}\n"
            f"게재일: {_fmt_date(p.published)} · 카테고리: {', '.join(p.categories)}\n"
            f"요약 근거: {s.basis}"
        )
        for w in s.warnings:
            warn = doc.add_paragraph()
            r = warn.add_run(f"⚠️ {w}")
            r.bold = True

        doc.add_heading("핵심 기여", level=3)
        for c in s.key_contributions or ["(내용 없음)"]:
            doc.add_paragraph(c, style="List Bullet")

        doc.add_heading("방법론", level=3)
        doc.add_paragraph(s.methodology or "(내용 없음)")

        doc.add_heading("실무 적용 포인트 (에이전트 의견 — 논문 내용과 구분됨)", level=3)
        for pt in s.practical_points or ["(내용 없음)"]:
            doc.add_paragraph(pt, style="List Bullet")

        doc.add_heading("한계", level=3)
        for lim in s.limitations or ["논문에 명시된 한계가 확인되지 않음"]:
            doc.add_paragraph(lim, style="List Bullet")

    doc.add_heading("[추출 실패] 목록", level=1)
    if data.failed:
        for item in data.failed:
            note = " → 초록 기반 요약으로 대체" if item.fallback_used else " → 요약 제외"
            doc.add_paragraph(
                f"{item.paper.title} ({item.paper.arxiv_id}) — {item.reason}{note}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("없음", style="List Bullet")

    if data.skipped_duplicates:
        doc.add_heading("중복 제외 (아카이브에 기존 요약 존재)", level=1)
        for p in data.skipped_duplicates:
            doc.add_paragraph(f"{p.title} ({p.arxiv_id})", style="List Bullet")

    notice = doc.add_paragraph()
    run = notice.add_run("\n" + APPROVAL_NOTICE)
    run.italic = True
    run.font.size = Pt(9)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


def briefing_paths(output_dir: Path, when: date | None = None) -> tuple[Path, Path]:
    """파일명 규칙 paper_brief_YYYYMMDD.docx 에 따른 (docx, md) 경로."""
    d = (when or date.today()).strftime("%Y%m%d")
    return output_dir / f"paper_brief_{d}.docx", output_dir / f"paper_brief_{d}.md"
